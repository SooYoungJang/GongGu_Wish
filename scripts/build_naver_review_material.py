from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "naver-login-service-introduction.docx"

ASSETS = {
    "home": ROOT / "apps" / "mobile" / "01-home-screen.png",
    "ranking": ROOT / "apps" / "mobile" / "store-assets" / "google-play" / "phone-01-ranking.png",
    "reels": ROOT / "apps" / "mobile" / "store-assets" / "google-play" / "phone-02-reels.png",
    "detail": ROOT / "apps" / "mobile" / "store-assets" / "google-play" / "phone-03-detail.png",
    "calendar": ROOT / "screenshots" / "gonggu-calendar-20260614-193459.png",
    "submit": ROOT / "screenshots" / "03-submit-screen.png",
    "auth": ROOT / "apps" / "mobile" / "evidence" / "gon-132-reverify-02-auth-login.png",
}

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RED = "9B1C1C"
GREEN = "0B6B42"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color=INK, bold=False):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, edge_data in kwargs.items():
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_paragraph(paragraph, before=0, after=6, line=1.10, alignment=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def create_numbering(doc, fmt, marker):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r_pr.append(r_fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_body(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=11 if style == "Normal" else 9, color=INK if style == "Normal" else MUTED)
    return paragraph


def add_bullet(doc, text, bullet_num_id):
    paragraph = add_body(doc, text)
    apply_numbering(paragraph, bullet_num_id)
    return paragraph


def add_numbered(doc, text, decimal_num_id):
    paragraph = add_body(doc, text)
    apply_numbering(paragraph, decimal_num_id)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    size, color, before, after = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }[level]
    set_run_font(run, size=size, color=color, bold=True)
    style_paragraph(paragraph, before=before, after=after, line=1.10)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def fill_cell(cell, text, size=9.5, color=INK, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, before=0, after=2, line=1.05, alignment=alignment)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_label_detail_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2700, 6660])
    for label, value in rows:
        cells = table.add_row().cells
        fill_cell(cells[0], label, bold=True, color=NAVY)
        fill_cell(cells[1], value)
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "D7DBE2"},
                bottom={"val": "single", "sz": 4, "color": "D7DBE2"},
                left={"val": "single", "sz": 4, "color": "D7DBE2"},
                right={"val": "single", "sz": 4, "color": "D7DBE2"},
            )
    return table


def add_callout(doc, title, text, fill=CALLOUT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 18, "color": accent},
        top={"val": "single", "sz": 4, "color": "D7DBE2"},
        bottom={"val": "single", "sz": 4, "color": "D7DBE2"},
        right={"val": "single", "sz": 4, "color": "D7DBE2"},
    )
    cell.text = ""
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, before=0, after=3, line=1.10)
    lead = paragraph.add_run(title)
    set_run_font(lead, size=10.5, color=accent, bold=True)
    body = paragraph.add_run("\n" + text)
    set_run_font(body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path, caption, width=5.6):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=3, after=3, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    style_paragraph(cap, before=2, after=5, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_two_figures(doc, left_path, left_caption, right_path, right_caption):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680], indent=0)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    for cell, path, caption in zip(table.rows[0].cells, [left_path, right_path], [left_caption, right_caption]):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, before=0, after=2, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.add_run().add_picture(str(path), width=Inches(2.85))
        cap = cell.add_paragraph()
        style_paragraph(cap, before=2, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = cap.add_run(caption)
        set_run_font(run, size=8.5, color=MUTED, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in [(1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)]:
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, size=9, color=MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.line_spacing = 1.0

    kicker = doc.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(kicker, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(6)
    kicker.paragraph_format.line_spacing = 1.0


def set_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    style_paragraph(hp, before=0, after=0, line=1.0)
    run = hp.add_run("GongGu Wish  |  네이버 로그인 검수 소명자료")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    style_paragraph(fp, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = fp.add_run("GongGu Wish  |  Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(fp)


def add_url_row(table, label, url, description):
    cells = table.add_row().cells
    fill_cell(cells[0], label, size=9.5, bold=True, color=NAVY)
    cells[1].text = ""
    paragraph = cells[1].paragraphs[0]
    style_paragraph(paragraph, before=0, after=2, line=1.05)
    add_hyperlink(paragraph, url, url)
    desc = paragraph.add_run("\n" + description)
    set_run_font(desc, size=9.3, color=INK)
    set_cell_shading(cells[0], LIGHT_BLUE)
    for cell in cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": "D7DBE2"},
            bottom={"val": "single", "sz": 4, "color": "D7DBE2"},
            left={"val": "single", "sz": 4, "color": "D7DBE2"},
            right={"val": "single", "sz": 4, "color": "D7DBE2"},
        )


def add_bullet_to_cell(cell, text, bullet_num_id):
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, before=0, after=5, line=1.10)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, color=INK)
    apply_numbering(paragraph, bullet_num_id)


def build():
    for path in ASSETS.values():
        if not path.exists():
            raise FileNotFoundError(path)

    doc = Document()
    set_document_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_header_footer(section)

    bullet_id = create_numbering(doc, "bullet", "•")
    decimal_id = create_numbering(doc, "decimal", "%1.")

    # Page 1 - cover and rejection response map.
    kicker = doc.add_paragraph(style="Kicker")
    run = kicker.add_run("NAVER LOGIN REVIEW  /  SERVICE EXPLANATION")
    set_run_font(run, size=10, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    style_paragraph(title, before=0, after=5, line=1.0)
    run = title.add_run("네이버 로그인 검수 소명자료")
    set_run_font(run, size=28, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, before=0, after=17, line=1.15)
    run = subtitle.add_run("GongGu Wish (공구위시) 서비스 소개 및 서비스 URL 안내")
    set_run_font(run, size=14, color=MUTED)
    add_label_detail_table(doc, [
        ("서비스명", "공구위시 (GongGu Wish)"),
        ("콘텐츠 유형", "인플루언서 공동구매 정보 큐레이션·일정 관리·알림 서비스"),
        ("운영 표면", "iOS/Android 모바일 앱 + 공개 웹 서비스"),
        ("작성일", "2026년 8월 10일"),
        ("자료 목적", "네이버 로그인 적용 서비스의 업종, 콘텐츠, 메뉴별 기능, 로그인 사용 목적 소명"),
    ])
    add_callout(doc, "검수자 확인 포인트", "공구위시는 상품을 직접 판매하는 쇼핑몰이 아니라, 공개 게시물에서 확인한 인플루언서 공동구매 정보를 검수·정리하여 제공하는 정보 서비스입니다. 네이버 로그인은 공개 콘텐츠를 보기 위한 필수 로그인보다, 북마크·팔로우·알림·계정 관리 등 개인화 기능을 위한 계정 인증에 사용됩니다.", fill=LIGHT_BLUE, accent=BLUE)
    add_heading(doc, "반려 사유 대응 요약", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2150, 3850, 3360])
    for cell, header in zip(table.rows[0].cells, ["검수 요청 항목", "자료에 포함한 내용", "확인 위치"]):
        fill_cell(cell, header, size=9.5, color=NAVY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, LIGHT_GRAY)
    for row in [
        ("서비스 콘텐츠 확인", "공동구매 상품·판매자·기간·할인·구매 링크를 제공하는 서비스임을 설명", "2~3쪽"),
        ("메뉴별 상세 설명 및 화면", "홈·랭킹·릴스·상세·캘린더·제보·로그인 화면과 기능을 정리", "4~8쪽"),
        ("전자상거래 여부", "직접 결제·배송 없이 외부 판매처로 연결하는 정보 서비스임을 명시", "3쪽"),
    ]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            fill_cell(cell, value, size=9.2)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val": "single", "sz": 4, "color": "D7DBE2"}, bottom={"val": "single", "sz": 4, "color": "D7DBE2"}, left={"val": "single", "sz": 4, "color": "D7DBE2"}, right={"val": "single", "sz": 4, "color": "D7DBE2"})

    # Page 2 - service summary and URLs.
    doc.add_page_break()
    add_heading(doc, "1. 서비스 개요", 1)
    add_body(doc, "공구위시(GongGu Wish)는 인스타그램 등 공개 게시물에 흩어진 인플루언서 공동구매 정보를 모아 이용자가 일정과 상품 정보를 한눈에 확인할 수 있도록 제공하는 모바일 중심 서비스입니다.")
    add_body(doc, "공개 게시물에서 상품명, 브랜드명, 카테고리, 공동구매 시작일·마감일, 할인 정보, 인플루언서 정보, 원본 게시물 및 외부 구매 링크를 확인하고, 운영자 검수 후 승인된 정보만 사용자 화면에 노출합니다.")
    add_heading(doc, "서비스 콘텐츠의 흐름", 2)
    for text in [
        "공개 Instagram 게시물 또는 이용자 제보에서 공동구매 후보 정보를 수집합니다.",
        "상품명·기간·할인·구매 링크 등 콘텐츠를 구조화하고 운영자가 공동구매 여부와 정보 품질을 확인합니다.",
        "승인된 공구를 홈, 랭킹, 릴스, 검색, 캘린더, 상품 상세 화면에 제공하고, 사용자가 외부 구매 페이지를 확인할 수 있도록 연결합니다.",
    ]:
        add_numbered(doc, text, decimal_id)
    add_heading(doc, "2. 접근 가능한 서비스 URL", 1)
    url_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(url_table, [2700, 6660])
    add_url_row(url_table, "서비스 소개", "https://gongguwish.com/", "공구위시 서비스의 업종과 핵심 기능을 소개하는 공개 웹 페이지")
    add_url_row(url_table, "공동구매 캘린더", "https://gongguwish.com/calendar", "날짜별 공동구매 콘텐츠를 확인하는 공개 페이지")
    add_url_row(url_table, "공동구매 제보", "https://gongguwish.com/submit", "발견한 공동구매 정보를 운영자에게 제보하는 공개 페이지")
    add_url_row(url_table, "이용약관", "https://gongguwish.com/terms", "공동구매 탐색·캘린더·제보·계정 기능의 이용 조건")
    add_url_row(url_table, "개인정보처리방침", "https://gongguwish.com/privacy", "앱과 웹 서비스의 개인정보 처리 안내")
    add_url_row(url_table, "계정 삭제 안내", "https://gongguwish.com/account-deletion", "서비스 계정 삭제 절차 안내")
    add_callout(doc, "모바일 앱 접근 안내", "네이버 로그인이 적용된 화면은 공구위시 모바일 앱에서 확인합니다. 앱 실행 후 공개 콘텐츠를 탐색할 수 있으며, 마이페이지 → 로그인 → 네이버로 계속하기 순서로 로그인 화면에 접근합니다. 공개 웹 URL은 서비스 콘텐츠 유형과 웹에서 제공되는 캘린더·제보 기능을 확인하기 위한 접근 경로입니다.", fill=CALLOUT, accent=GREEN)

    # Page 3 - industry, commerce boundary, login purpose.
    doc.add_page_break()
    add_heading(doc, "3. 서비스 업종 및 콘텐츠 형태", 1)
    add_label_detail_table(doc, [
        ("서비스 업종", "콘텐츠 정보 서비스 / 공동구매 일정·검색·알림·큐레이션"),
        ("주요 이용자", "인플루언서 공동구매 상품과 오픈·마감 일정을 빠르게 확인하려는 이용자"),
        ("콘텐츠 형태", "상품 카드, 인플루언서·판매자 정보, 이미지·영상, 일정, 할인 정보, 원본 게시물 링크, 외부 구매 링크"),
        ("상품 범위", "식품, 생활용품, 뷰티, 패션, 홈·리빙, 육아 등 인플루언서 공동구매에서 다뤄지는 상품 정보"),
        ("운영 방식", "공개 콘텐츠 자동 수집·정보 보강 + 이용자 제보 + 운영자 검수·승인"),
    ])
    add_heading(doc, "전자상거래 관련 소명", 2)
    add_callout(doc, "직접 판매·결제 서비스가 아님", "공구위시는 상품을 직접 매입하거나 판매하지 않습니다. 앱과 웹에서는 공동구매 상품의 정보와 일정, 판매자, 할인 정보 및 외부 구매 링크를 제공합니다. 구매 링크를 누르면 각 인플루언서 또는 외부 판매처의 구매 페이지로 이동하며, 공구위시 안에서는 주문, 결제, 재고, 배송, 환불을 처리하지 않습니다.", fill="FFF8E8", accent="7A5A00")
    add_body(doc, "따라서 본 서비스의 핵심 콘텐츠는 전자상거래 거래 자체가 아니라, 여러 공개 출처에 흩어진 공동구매 정보를 일정과 탐색 화면으로 정리해 제공하는 큐레이션 콘텐츠입니다.")
    add_heading(doc, "4. 네이버 로그인 적용 목적", 1)
    add_body(doc, "공개 콘텐츠는 로그인 없이 확인할 수 있습니다. 네이버 로그인은 서비스 계정을 생성·식별하고 다음 개인화 기능을 이용하기 위한 선택적 인증 수단입니다.")
    for text in [
        "관심 공동구매 북마크 및 마이페이지에서 저장 목록 확인",
        "관심 인플루언서·판매자 팔로우 및 관련 콘텐츠 탐색",
        "공구 시작·마감 알림과 알림 설정 관리",
        "최근 본 공구, 계정 설정, 개인정보·회원탈퇴 등 계정 관리",
    ]:
        add_bullet(doc, text, bullet_id)
    add_heading(doc, "네이버 로그인 이용 순서", 2)
    for text in [
        "앱 실행 후 홈·랭킹·릴스·검색 등 공개 콘텐츠를 확인합니다.",
        "마이페이지에서 로그인 화면을 엽니다.",
        "로그인 화면의 네이버로 계속하기를 선택합니다.",
        "네이버 인증을 완료하면 공구위시 앱으로 돌아와 마이페이지와 개인화 기능을 이용합니다.",
    ]:
        add_numbered(doc, text, decimal_id)
    add_callout(doc, "로그인 기능의 범위", "네이버 로그인은 공구 콘텐츠를 네이버에서 검색하거나 네이버 계정 정보를 서비스 콘텐츠로 게시하기 위한 기능이 아닙니다. 공구위시 서비스 내부의 계정 인증과 개인화 기능 제공을 위한 로그인입니다.", fill=LIGHT_BLUE, accent=BLUE)

    # Page 4 - menu details.
    doc.add_page_break()
    add_heading(doc, "5. 메뉴별 상세 기능", 1)
    add_body(doc, "아래 표는 네이버 로그인이 적용되는 모바일 앱의 주요 메뉴와, 각 메뉴에서 제공하는 콘텐츠·기능을 정리한 것입니다.")
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1400, 2500, 3700, 1760])
    for cell, header in zip(table.rows[0].cells, ["메뉴", "콘텐츠", "주요 기능", "로그인과의 관계"]):
        fill_cell(cell, header, size=9.2, color=NAVY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, LIGHT_GRAY)
    menu_rows = [
        ("홈", "이번주·추천 공구, 프로모션, 카테고리", "공구 카드 확인, 카테고리 필터, 검색 진입, 날짜별 공구 이동", "공개 탐색 가능; 저장·알림은 계정 기능"),
        ("랭킹", "인기 공구·인플루언서/판매자 순위", "오늘·이번주·이번달, 카테고리, 인기 셀러·인기 공구·신규 오픈·마감 임박 필터", "공개 열람 가능; 팔로우 상태 저장에 로그인 사용"),
        ("릴스", "공구 상품 이미지·영상 기반 세로 피드", "영상/이미지 탐색, 상품 요약, 판매자, 가격·마감, 북마크·알림·구매 링크", "공개 열람 가능; 개인화 액션에 로그인 사용"),
        ("검색", "브랜드명·제품명·셀러명과 공구 데이터", "키워드 검색, 인기 검색어, 카테고리 및 결과 탐색", "공개 검색 가능"),
        ("캘린더", "날짜별 승인 공구 일정", "월간 날짜 이동, 오늘 이동, 공구 목록, 북마크·알림 필터, 상세 이동", "공개 일정 열람; 개인 활동 필터에 로그인/저장 상태 사용"),
        ("상품 상세", "상품·브랜드·판매자·기간·할인·요약·원본/구매 링크", "상세 정보 확인, 북마크, 알림, 공유, 외부 구매 페이지 이동", "공개 상세 확인; 저장·알림에 로그인 사용"),
        ("공구 제보", "Instagram 원본 URL과 상품·기간·구매 정보", "URL 입력, 콘텐츠 정보 자동 보강, 추가 정보 수정, 운영자 검수 제출", "공개 제보 경로 제공; 제출 정책에 따라 계정 확인이 적용될 수 있음"),
        ("마이", "프로필과 사용자 활동", "북마크, 최근 본 공구, 알림 공구, 팔로우, 설정, 계정 삭제", "네이버 로그인 적용 핵심 메뉴"),
    ]
    for row in menu_rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            fill_cell(cell, value, size=8.8)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val": "single", "sz": 4, "color": "D7DBE2"}, bottom={"val": "single", "sz": 4, "color": "D7DBE2"}, left={"val": "single", "sz": 4, "color": "D7DBE2"}, right={"val": "single", "sz": 4, "color": "D7DBE2"})
    add_body(doc, "※ 화면 캡처의 상품명·판매자명·수량은 테스트 또는 스토어 소개용 샘플일 수 있으며, 실제 공개 데이터는 공구 진행 상태와 검수 결과에 따라 달라질 수 있습니다.", style="Caption")

    # Page 5 - home and ranking.
    doc.add_page_break()
    add_heading(doc, "6. 화면 예시 - 홈과 랭킹", 1)
    add_body(doc, "홈은 서비스의 핵심 콘텐츠인 공동구매 정보를 처음 만나는 화면입니다. 랭킹은 기간·카테고리별로 인기 공구와 판매자를 탐색하는 화면입니다.")
    add_figure(doc, ASSETS["home"], "그림 1. 홈/공구 탐색 화면 예시", width=4.9)
    add_body(doc, "홈 화면에서는 인기 공구 카드, 카테고리와 일정 정보를 확인할 수 있습니다. 공구 카드에는 상품명과 가격·마감 정보가 노출되어 서비스 콘텐츠 유형을 즉시 파악할 수 있습니다.")
    add_figure(doc, ASSETS["ranking"], "그림 2. 랭킹/공구 탐색 화면의 상품·기간·판매자 정보 예시", width=4.9)
    add_callout(doc, "검수 포인트", "서비스의 중심 콘텐츠는 상품 자체의 판매 페이지가 아니라, 공동구매 상품 정보와 진행 기간을 모아 보여주는 정보 카드입니다.", fill=CALLOUT, accent=BLUE)

    # Page 6 - reels and detail.
    doc.add_page_break()
    add_heading(doc, "7. 화면 예시 - 릴스와 상품 상세", 1)
    add_body(doc, "릴스는 공동구매 상품의 이미지·영상 콘텐츠를 세로로 빠르게 탐색하는 화면입니다. 상품 상세에서는 해당 공구의 기간, 가격, 판매자, 카테고리, 요약과 외부 구매 링크를 확인합니다.")
    add_two_figures(doc, ASSETS["reels"], "그림 3. 릴스형 공동구매 콘텐츠 탐색", ASSETS["detail"], "그림 4. 상품 상세의 상품·판매자·구매 링크 정보")
    add_body(doc, "상품 상세의 구매 링크는 각 공구 판매자 또는 외부 쇼핑몰이 제공하는 페이지로 이동하기 위한 링크입니다. 공구위시의 역할은 해당 링크를 포함한 공구 정보를 검수·정리하여 탐색 가능하게 제공하는 것입니다.")
    add_callout(doc, "상품 취급 범위", "현재 서비스에서 소개하는 상품은 인플루언서 공동구매 콘텐츠에 포함된 식품, 생활용품, 뷰티, 패션, 홈·리빙, 육아 등의 상품 정보입니다. 공구위시가 직접 상품을 보유하거나 판매하지는 않습니다.", fill="FFF8E8", accent="7A5A00")

    # Page 7 - calendar and submission.
    doc.add_page_break()
    add_heading(doc, "8. 화면 예시 - 공개 웹 캘린더와 공구 제보", 1)
    add_body(doc, "공개 웹은 모바일 앱과 동일한 서비스의 콘텐츠 확인용 접근 경로입니다. 캘린더 페이지에서는 승인된 공동구매의 상품명, 판매자, 요약, 할인 및 마감 시점을 확인할 수 있습니다.")
    add_figure(doc, ASSETS["calendar"], "그림 5. 공개 웹 공동구매 캘린더 화면 예시", width=5.8)
    add_body(doc, "공구 제보 화면에서는 Instagram 원본 게시물 URL을 입력하고 제품명, 카테고리, 시작일·종료일, 구매 링크, 할인 정보, 요약을 제출합니다. 게시물 정보가 확인되면 상품 정보 자동 보강을 활용할 수 있으며, 제출된 정보는 운영자 검수 후 서비스에 반영됩니다.")
    add_figure(doc, ASSETS["submit"], "그림 6. 공동구매 제보 입력 화면 예시", width=3.4)

    # Page 8 - auth and account.
    doc.add_page_break()
    add_heading(doc, "9. 화면 예시 - 로그인과 마이페이지", 1)
    add_body(doc, "네이버 로그인은 마이페이지의 개인화 기능을 이용하기 위한 인증 단계입니다. 공개 화면을 먼저 탐색한 뒤, 저장·팔로우·알림·계정 관리가 필요한 경우 로그인 화면으로 이동합니다.")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4200, 5160], indent=0)
    left, right = table.rows[0].cells
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    left.text = ""
    paragraph = left.paragraphs[0]
    style_paragraph(paragraph, before=0, after=2, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.add_run().add_picture(str(ASSETS["auth"]), width=Inches(2.65))
    caption = left.add_paragraph()
    style_paragraph(caption, before=2, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = caption.add_run("그림 7. 로그인 화면 예시(공급자 구성은 빌드별 차이 가능)")
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    right.text = ""
    paragraph = right.paragraphs[0]
    style_paragraph(paragraph, before=2, after=6, line=1.10)
    run = paragraph.add_run("현재 구현 기준: 네이버로 계속하기")
    set_run_font(run, size=12, color=GREEN, bold=True)
    for text in [
        "로그인 화면에서 ‘네이버로 계속하기’를 선택합니다.",
        "인증 완료 후 공구위시 앱으로 복귀합니다.",
        "마이페이지에서 북마크·최근 본 공구·알림 공구·팔로우·설정을 확인합니다.",
        "계정 삭제 안내와 개인정보처리방침 링크를 제공합니다.",
    ]:
        add_bullet_to_cell(right, text, bullet_id)
    add_callout(doc, "로그인 없이 가능한 이용", "홈·랭킹·릴스·검색·캘린더·상품 상세 같은 공개 콘텐츠는 로그인 없이 확인할 수 있습니다. 네이버 로그인은 개인 활동을 저장하고 계정별 설정을 관리하기 위한 기능입니다.", fill=LIGHT_BLUE, accent=GREEN)
    doc.add_page_break()
    add_heading(doc, "10. 검수자가 확인할 수 있는 서비스 예정 URL", 1)
    add_body(doc, "웹 서비스는 현재 위 공개 URL을 통해 콘텐츠 확인이 가능하며, 모바일 앱은 같은 공구위시 서비스의 iOS/Android 표면으로 제공합니다. 모바일 앱에서의 네이버 로그인 적용 화면은 다음 경로로 확인할 수 있습니다.")
    for text in [
        "공구위시 앱 실행 → 마이페이지",
        "로그인 또는 개인화 기능 선택",
        "네이버로 계속하기 선택",
        "인증 후 마이페이지로 복귀",
    ]:
        add_numbered(doc, text, decimal_id)
    add_body(doc, "앱 검수용 설치 경로 또는 테스터 접근 정보가 필요한 경우에는 네이버 로그인 검수 신청 페이지의 앱 접근 정보란에 해당 빌드/테스트 계정을 함께 기재해 주시면 됩니다.")

    # Page 9 - copy text appendix.
    doc.add_page_break()
    add_heading(doc, "부록 A. 검수 신청 페이지 소명 내용", 1)
    add_body(doc, "아래 문안은 네이버 로그인 검수 신청 페이지의 ‘소명 내용 입력’ 란에 그대로 붙여 넣을 수 있도록 작성했습니다.")
    submission = (ROOT / "output" / "naver-login-review" / "naver-login-submission-explanation.txt").read_text(encoding="utf-8")
    add_callout(doc, "복사하여 제출할 문안", submission, fill=CALLOUT, accent=BLUE)
    add_heading(doc, "부록 B. 자료 작성 기준", 1)
    add_body(doc, "본 자료는 공구위시의 모바일 앱 화면 구조, 공개 웹 서비스 페이지, 서비스 이용약관·개인정보처리방침에 기재된 서비스 기능을 기준으로 작성했습니다. 화면 캡처에는 테스트 또는 스토어 소개용 샘플 데이터가 포함될 수 있으며, 서비스 운영 데이터는 검수·기간·콘텐츠 상태에 따라 변경될 수 있습니다.")
    add_body(doc, "검수자가 접근할 URL: https://gongguwish.com/  |  캘린더: https://gongguwish.com/calendar  |  제보: https://gongguwish.com/submit", style="Caption")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
